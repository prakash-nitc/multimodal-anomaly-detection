# -*- coding: utf-8 -*-
"""Convert the project handbook from Markdown to a formatted Word document.

Deliberately not a general Markdown converter. It handles exactly the
constructs the handbook uses -- headings, tables, blockquotes, lists, inline
emphasis and code, horizontal rules -- and formats them for a document that will
be read end to end and probably printed, rather than for fidelity to some
Markdown specification.

Typography follows the rest of the project's materials: Segoe UI for headings
and tables, Georgia for running text since the document is long-form reading,
and the same teal accent used in the slide deck and the published page.
"""
from __future__ import annotations

import os
import re
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

INK = RGBColor(0x14, 0x1C, 0x1E)
INK_2 = RGBColor(0x3A, 0x4A, 0x4C)
MUTED = RGBColor(0x61, 0x72, 0x74)
ACCENT = RGBColor(0x0D, 0x6B, 0x67)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

SANS = "Segoe UI"
SERIF = "Georgia"
MONO = "Consolas"

SHADE_HEAD = "141C1E"
SHADE_ALT = "F5F7F7"
SHADE_HI = "DDEDEB"

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*\s][^*]*?\*)")


def shade(cell, hexcolor: str) -> None:
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(el)


def bar(paragraph, hexcolor: str, size: int = 18) -> None:
    """A coloured left border -- used for blockquotes."""
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(size))
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), hexcolor)
    borders.append(left)
    pPr.append(borders)


def rule(paragraph, hexcolor: str = "D8E0DF") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hexcolor)
    borders.append(bottom)
    pPr.append(borders)


def add_runs(paragraph, text: str, size: float, font: str = SERIF,
             color=INK_2, bold_all: bool = False) -> None:
    """Emit text into a paragraph, honouring **bold**, *italic* and `code`."""
    for part in INLINE.split(text):
        if not part:
            continue
        b, i, mono = bold_all, False, False
        if part.startswith("**") and part.endswith("**"):
            part, b = part[2:-2], True
        elif part.startswith("`") and part.endswith("`"):
            part, mono = part[1:-1], True
        elif part.startswith("*") and part.endswith("*"):
            part, i = part[1:-1], True
        r = paragraph.add_run(part)
        r.font.size = Pt(size - (0.5 if mono else 0))
        r.font.name = MONO if mono else font
        r.font.bold = b
        r.font.italic = i
        r.font.color.rgb = INK if (b and not bold_all) else color


def heading(doc, text: str, level: int, first: bool) -> None:
    if level == 1 and not first:
        doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2 if level == 1 else (16 if level == 2 else 12))
    pf.space_after = Pt(8 if level == 1 else 6)
    pf.keep_with_next = True
    size = {1: 20, 2: 15, 3: 12.5, 4: 11}[level]
    r = p.add_run(re.sub(r"[*`]", "", text))
    r.font.size = Pt(size)
    r.font.name = SANS
    r.font.bold = True
    r.font.color.rgb = ACCENT if level <= 2 else INK
    if level == 1:
        rule(p, "0D6B67")


def add_table(doc, rows: list[list[str]]) -> None:
    ncol = max(len(r) for r in rows)
    rows = [r + [""] * (ncol - len(r)) for r in rows]
    t = doc.add_table(rows=len(rows), cols=ncol)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            hi = "**" in val and ri > 0
            if ri == 0:
                shade(cell, SHADE_HEAD)
                add_runs(p, val.replace("**", ""), 8.5, SANS, WHITE, bold_all=True)
            else:
                shade(cell, SHADE_HI if hi else (SHADE_ALT if ri % 2 else "FFFFFF"))
                add_runs(p, val, 9, SANS, INK_2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def convert(md_path: str, out_path: str) -> None:
    lines = open(md_path, encoding="utf-8").read().split("\n")

    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.27), Inches(11.69)   # A4
    sec.left_margin = sec.right_margin = Inches(0.9)
    sec.top_margin = Inches(0.85)
    sec.bottom_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = SERIF
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.22

    seen_h1 = False
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()

        if not s:
            i += 1
            continue

        # --- tables -------------------------------------------------
        if s.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells if c):
                    block.append(cells)
                i += 1
            if block:
                add_table(doc, block)
            continue

        # --- blockquotes --------------------------------------------
        if s.startswith(">"):
            block = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                block.append(lines[i].strip().lstrip(">").strip())
                i += 1
            text = " ".join(x for x in block if x)
            for chunk in [c for c in re.split(r"\s{2,}", text) if c]:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.22)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                bar(p, "0D6B67")
                add_runs(p, chunk, 10.5, SERIF, INK)
            continue

        # --- horizontal rule ----------------------------------------
        if re.fullmatch(r"-{3,}", s):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(10)
            rule(p)
            i += 1
            continue

        # --- headings -----------------------------------------------
        m = re.match(r"(#{1,4})\s+(.*)", s)
        if m:
            lvl = len(m.group(1))
            heading(doc, m.group(2), lvl, not seen_h1)
            if lvl == 1:
                seen_h1 = True
            i += 1
            continue

        # --- lists ---------------------------------------------------
        m = re.match(r"[-*]\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, m.group(1), 10.5)
            i += 1
            continue
        m = re.match(r"(\d+)\.\s+(.*)", s)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.space_after = Pt(3)
            add_runs(p, m.group(2), 10.5)
            i += 1
            continue

        # --- italic-only line (the dateline) -------------------------
        if s.startswith("*") and s.endswith("*") and not s.startswith("**"):
            p = doc.add_paragraph()
            r = p.add_run(s.strip("*"))
            r.font.size = Pt(9.5)
            r.font.italic = True
            r.font.name = SANS
            r.font.color.rgb = MUTED
            i += 1
            continue

        # --- ordinary paragraph, joining wrapped lines ---------------
        block = []
        while i < len(lines):
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("|", ">", "#"))
                    or re.match(r"[-*]\s+", nxt) or re.match(r"\d+\.\s+", nxt)
                    or re.fullmatch(r"-{3,}", nxt)):
                break
            block.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_runs(p, " ".join(block), 10.5)

    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else (
        r"p:\Research\multimodal-anomaly-detection\docs\08_understanding\00_HANDBOOK.md")
    dst = sys.argv[2] if len(sys.argv) > 2 else (
        r"p:\Research\multimodal-anomaly-detection\docs\08_understanding\DA-ZVAD_Handbook.docx")
    convert(src, dst)
    print(f"saved: {dst}  ({os.path.getsize(dst) / 1024:.0f} KB)")
